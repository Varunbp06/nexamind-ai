import hashlib
from io import BytesIO
import re
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from loguru import logger
from pairag.file.readers.base import BaseReader, FileItem, Document, List
from pairag.file.store.base import BaseFileStore
from pairag.file.utils.image_utils import compress_image_if_needed
from pairag.file.utils.markdown_tree_utils import (
    PaiTable,
    convert_table_to_markdown,
    is_horizontal_table,
)
from pairag.file.utils.image_caption_tool import ImageCaptionTool
from pairag.file.utils.image_utils import to_markdown_image_text
from pairag.file.utils.text_utils import replace_consecutive_spaces

class DocxReader(BaseReader):
    def __init__(
        self, file_store: BaseFileStore, image_caption_tool: ImageCaptionTool = None
    ):
        self.file_store = file_store
        self.image_caption_tool = image_caption_tool
        logger.info("DocxReader inited.")

    def _convert_paragraph(self, paragraph):
        text = paragraph.text.strip()
        if not text:
            return ""

        # 处理标题
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            heading_level = int(
                re.search(r"Heading (\d)", paragraph.style.name).group(1)
            )
            if heading_level > 6:
                heading_level = 6
            return f"{'#' * heading_level} {text}\n\n"

        # 处理普通段落
        return f"{text}\n\n"

    def _get_list_level(self, paragraph):
        indent_levels = {
            "List Paragraph": 0,
            "List Bullet": 1,
            "List Number": 1,
            "List Bullet 2": 2,
            "List Number 2": 2,
            "List Bullet 3": 3,
            "List Number 3": 3,
        }

        # 获取段落的样式名称
        style_name = paragraph.style.name
        # 根据样式名称获取层级
        return indent_levels.get(style_name, 0)

    def _convert_list(self, paragraph, level=0, number=1):
        text = paragraph.text.strip()
        if not text:
            return ""

        # 处理无序列表
        if paragraph.style.name.startswith("List Bullet") or not self._is_ordered_list(paragraph):
            indent = "  " * level
            return f"{indent}- {text}\n"
        # 处理有序列表
        elif paragraph.style.name.startswith("List Number") or paragraph.style.name.startswith("List Paragraph"):
            indent = "  " * level
            return f"{indent}{number}. {text}\n"

        return ""

    def _convert_table_to_markdown(self, table, doc_name):
        total_cols = max(len(row.cells) for row in table.rows)

        table_matrix = []
        for row in table.rows:
            table_matrix.append(self._parse_row(row, doc_name, total_cols))
        if is_horizontal_table(table_matrix):
            table = PaiTable(data=table_matrix, row_headers_index=[0])
        else:
            table = PaiTable(data=table_matrix, column_headers_index=[0])
        return convert_table_to_markdown(table, total_cols)

    def _parse_row(self, row, doc_name, total_cols):
        row_cells = [""] * total_cols
        col_index = 0
        for cell in row.cells:
            while col_index < total_cols and row_cells[col_index] != "":
                col_index += 1
            if col_index >= total_cols:
                break
            cell_content = self._parse_cell(cell, doc_name).strip()
            row_cells[col_index] = cell_content
            col_index += 1
        return row_cells

    def _parse_cell(self, cell, doc_name):
        cell_content = []
        for paragraph in cell.paragraphs:
            parsed_paragraph = self._parse_cell_paragraph(paragraph, doc_name)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content) if unique_content else ""

    def _parse_cell_paragraph(self, paragraph, doc_name):
        # 使用paragraph.text直接获取所有文本内容，这比遍历runs更可靠
        # paragraph.text会自动处理所有runs，包括格式化文本
        paragraph_text = paragraph.text.strip() if paragraph.text else ""
        
        run_texts = []
        for run in paragraph.runs:
            run_text = run.text if run.text is not None else ""
            run_texts.append(run_text)
        
        # 如果paragraph.text为空但runs有文本，使用runs的文本（处理特殊情况）
        if not paragraph_text and any(run_texts):
            paragraph_text = "".join(run_texts).strip()
        
        return paragraph_text
    

    def _is_ordered_list(self, paragraph) -> bool:
        """
        判断一个 paragraph 是否属于有序列表（如 1, 2, 3...），而不是无序列表（•, -）。
        返回 True 表示有序，False 表示无序或非列表。
        """
        p = paragraph._element
        num_pr = p.find('.//w:numPr', p.nsmap)
        if num_pr is None:
            return False

        num_id = num_pr.find('.//w:numId', p.nsmap)
        if num_id is None:
            return False

        # 获取 numbering part
        numbering_part = paragraph.part.numbering_part
        if numbering_part is None:
            return False

        # 查找对应的 num -> abstractNumId
        num_id_val = num_id.get(qn('w:val'))
        for num in numbering_part.element.findall('.//w:num', numbering_part.element.nsmap):
            if num.get(qn('w:numId')) == num_id_val:
                abstract_num_id = num.find('.//w:abstractNumId', num.nsmap)
                if abstract_num_id is None:
                    return False
                abstract_id_val = abstract_num_id.get(qn('w:val'))
                break
        else:
            return False
        
        # Numbered formats include: decimal, lowerRoman, upperRoman, lowerLetter, upperLetter
        # Bullet formats include: bullet
        numbered_formats = {
            "decimal",
            "lowerRoman",
            "upperRoman",
            "lowerLetter",
            "upperLetter",
            "decimalZero",
        }

        # 查找 abstractNum 定义
        for abstract in numbering_part.element.findall('.//w:abstractNum', numbering_part.element.nsmap):
            if abstract.get(qn('w:abstractNumId')) == abstract_id_val:
                ilvl_elem = num_pr.find('.//w:ilvl', p.nsmap)
                ilvl = ilvl_elem.get(qn('w:val')) if ilvl_elem is not None else '0'
                # 找到对应级别的 lvl
                for lvl in abstract.findall('.//w:lvl', abstract.nsmap):
                    if lvl.get(qn('w:ilvl')) == ilvl:
                        num_fmt = lvl.find('.//w:numFmt', lvl.nsmap)
                        if num_fmt is not None:
                            fmt = num_fmt.get(qn('w:val'))
                            return fmt in numbered_formats
                return False
        return False

       

    def convert_docx_to_markdown(
        self, document: DocxDocument, save_name_template: str, tenant_id: str,
    ) -> str:
        paragraphs = document.paragraphs.copy()
        tables = document.tables.copy()
        markdown = []
        images = []
        # 用于跟踪有序列表的编号
        ordered_counters = []
        # 记录当前列表层级（用于判断是否连续）
        last_list_level = -1
        for element in document.element.body:
            if isinstance(element.tag, str) and element.tag.endswith("p"):  # 段落
                paragraph = paragraphs.pop(0)

                if paragraph.style and paragraph.style.name.startswith("List"):
                    current_level = self._get_list_level(paragraph)
                    if current_level > last_list_level:
                        # 进入更深层级：压入新计数器
                        ordered_counters.append(1)
                    elif current_level < last_list_level:
                        # 退回上层：弹出多余层级
                        ordered_counters = ordered_counters[:current_level + 1]
                        if current_level >= 0:
                            ordered_counters[current_level] += 1
                        else:
                            ordered_counters = [1]
                    else:
                        # 同一层级：递增
                        if current_level < len(ordered_counters):
                            ordered_counters[current_level] += 1
                        else:
                            ordered_counters.append(1)

                    last_list_level = current_level

                    # 获取当前层级的编号
                    current_number = ordered_counters[current_level] if current_level < len(ordered_counters) else 1
                    markdown.append(self._convert_list(paragraph, current_level, current_number))
                else:
                    ordered_counters = []
                    last_list_level = -1
                    for run in paragraph.runs:
                        if (
                            hasattr(run.element, "tag")
                            and isinstance(element.tag, str)
                            and run.element.tag.endswith("r")
                        ):
                            drawing_elements = run.element.findall(
                                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                            )
                            for drawing in drawing_elements:
                                blip_elements = drawing.findall(
                                    ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                                )
                                for blip in blip_elements:
                                    embed_id = blip.get(
                                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                    )
                                    if (
                                        embed_id
                                        and self.image_caption_tool
                                    ):
                                        image_part = document.part.related_parts.get(
                                            embed_id
                                        )
                                        if hasattr(image_part, "blob"):
                                            image_blob = image_part.blob
                                            image_name = (
                                                hashlib.md5(image_blob).hexdigest()
                                                + ".jpeg"
                                            )

                                            save_image_name = save_name_template.format(
                                                image_name
                                            )
                                            image_file = BytesIO(image_blob)
                                            image_file = compress_image_if_needed(
                                                image_file
                                            )
                                            if not image_file:
                                                continue
                                            try:
                                                upload_result = self.file_store.write(
                                                    file=image_file,
                                                    file_name=image_name,
                                                    file_path=save_image_name,
                                                    tenant_id=tenant_id,
                                                )
                                                image_alt_text = self.image_caption_tool.extract_image(image_blob)
                                                if image_alt_text:
                                                    cleaned_alt = re.sub(r'\n', ' ', image_alt_text).replace('\r', '').strip()
                                                    image_text = to_markdown_image_text(upload_result.file_path, cleaned_alt)
                                                    markdown.append(f"{image_text}\n")
                                                    images.append(upload_result.file_path)

                                                    logger.info(
                                                        f"Successfully saved image {upload_result.file_path}."
                                                    )
                                            except Exception as ex:
                                                logger.exception(
                                                    f"Failed to save image from URL: {upload_result.file_path}. Error: {ex}"
                                                )

                    markdown.append(self._convert_paragraph(paragraph))

            elif isinstance(element.tag, str) and element.tag.endswith("tbl"):  # 表格
                table = tables.pop(0)
                markdown.append(self._convert_table_to_markdown(table, None))
                markdown.append("\n\n")

        return "".join(markdown), images

    def read(self, file_item: FileItem) -> List[Document]:
        """
        Read a CSV file and return a list of Documents.
        """
        try:
            file_item.file.seek(0)
            docx_file = DocxDocument(file_item.file)

            markdown_content, images = self.convert_docx_to_markdown(
                docx_file, file_item.kb_id + "/images/{}", tenant_id=file_item.tenant_id,
            )
            markdown_content = replace_consecutive_spaces(markdown_content)
            metadata = file_item.metadata()

            docs = [Document(id_=file_item.id, text=markdown_content, metadata=metadata)]
            logger.info(f"Successfully read {file_item.file_name}.")

            return docs
        except Exception as e:
            logger.exception(e)
            return []
